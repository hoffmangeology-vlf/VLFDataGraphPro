import os
import csv
import io
import re
import json
import math
import threading
import urllib.parse
import urllib.request
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from datetime import datetime, timezone

import numpy as np
import pandas as pd
import matplotlib.dates as mdates
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk


APP_TITLE = "VLF DataGraph Pro"
APP_VERSION = "0.9 Alpha"
APP_AUTHOR = "Michal Hoffman"
APP_LICENSE = "Free for non-commercial use"
DEFAULT_MAX_POINTS = 30000
MAX_ALLOWED_POINTS = 500000

USGS_QUERY_URL = "https://earthquake.usgs.gov/fdsnws/event/1/query"
NOAA_GOES_XRAY_URL = "https://services.swpc.noaa.gov/json/goes/primary/xrays-7-day.json"


def reduce_minmax(x, y, max_points=DEFAULT_MAX_POINTS):
    x_arr = np.asarray(x)
    y_arr = np.asarray(y, dtype=float)
    n = len(y_arr)

    if n <= max_points or max_points < 100:
        return x_arr, y_arr

    bucket_count = max(1, max_points // 2)
    edges = np.linspace(0, n, bucket_count + 1, dtype=int)
    out_x, out_y = [], []

    for i in range(bucket_count):
        start, end = edges[i], edges[i + 1]
        if end <= start:
            continue

        part = y_arr[start:end]
        valid = np.isfinite(part)
        if not valid.any():
            continue

        valid_idx = np.flatnonzero(valid)
        vals = part[valid]
        imin = valid_idx[np.argmin(vals)]
        imax = valid_idx[np.argmax(vals)]

        for idx in sorted({start + int(imin), start + int(imax)}):
            out_x.append(x_arr[idx])
            out_y.append(y_arr[idx])

    return np.asarray(out_x), np.asarray(out_y)


def safe_float(value):
    try:
        return float(value)
    except (TypeError, ValueError):
        return np.nan


def flare_class(flux):
    if not np.isfinite(flux) or flux <= 0:
        return ""
    if flux >= 1e-4:
        return f"X{flux / 1e-4:.1f}"
    if flux >= 1e-5:
        return f"M{flux / 1e-5:.1f}"
    if flux >= 1e-6:
        return f"C{flux / 1e-6:.1f}"
    if flux >= 1e-7:
        return f"B{flux / 1e-7:.1f}"
    return f"A{flux / 1e-8:.1f}"


class ScrollableCheckFrame(ttk.Frame):
    def __init__(self, master, width=210):
        super().__init__(master)
        self.canvas = tk.Canvas(self, highlightthickness=0, width=width)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.inner = ttk.Frame(self.canvas)

        self.inner.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")),
        )
        self.window_id = self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.canvas.bind(
            "<Configure>",
            lambda e: self.canvas.itemconfigure(self.window_id, width=e.width),
        )
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.canvas.bind("<Enter>", self._bind_mousewheel)
        self.canvas.bind("<Leave>", self._unbind_mousewheel)

    def _bind_mousewheel(self, _event):
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)

    def _unbind_mousewheel(self, _event):
        self.canvas.unbind_all("<MouseWheel>")

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-event.delta / 120), "units")


class SessionTab:
    def __init__(self, app, notebook, file_path):
        self.app = app
        self.notebook = notebook
        self.file_path = file_path
        self.df = None
        self.numeric_columns = []
        self.channel_vars = {}
        self.x_full = None
        self.x_is_datetime = False
        self.last_selected_columns = []
        self.axes = []
        self.lines = []
        self.earthquakes = []
        self.flares = []
        self.daily_stats = None
        self.loading = False

        self.frame = ttk.Frame(notebook)
        self.notebook.add(self.frame, text=os.path.basename(file_path)[:28])
        self.notebook.select(self.frame)

        self.status_var = tk.StringVar(value="Načítavam...")
        self.cursor_var = tk.StringVar(value="Cursor: —")
        self.max_points_var = tk.StringVar(value=str(DEFAULT_MAX_POINTS))
        self.plot_mode_var = tk.StringVar(value="Single Plot")
        self.line_width_var = tk.StringVar(value="0.8")
        self.show_grid_var = tk.BooleanVar(value=True)
        self.dark_mode_var = tk.BooleanVar(value=False)
        self.anomaly_var = tk.BooleanVar(value=False)
        self.sigma_var = tk.StringVar(value="3.0")
        self.show_quakes_var = tk.BooleanVar(value=True)
        self.show_flares_var = tk.BooleanVar(value=True)

        self._build_ui()
        self._new_figure()
        self.load_async()

    def _build_ui(self):
        top = ttk.Frame(self.frame, padding=6)
        top.pack(fill=tk.X)

        ttk.Button(top, text="Plot", command=self.plot_selected).pack(side=tk.LEFT)
        ttk.Button(top, text="Compare Days", command=self.compare_days).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(top, text="Daily Statistics", command=self.show_daily_stats).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(top, text="Import USGS Earthquakes", command=self.import_earthquakes).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(top, text="Import NOAA GOES Solar Flares", command=self.import_flares).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(top, text="Export Report", command=self.export_protocol).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(top, text="Save Graph", command=self.save_graph).pack(side=tk.LEFT, padx=(5, 0))

        ttk.Label(top, text="Plot Mode:").pack(side=tk.LEFT, padx=(15, 3))
        ttk.Combobox(
            top,
            textvariable=self.plot_mode_var,
            values=["Single Plot", "Stacked Panels"],
            state="readonly",
            width=15,
        ).pack(side=tk.LEFT)

        ttk.Label(top, text="Max Points:").pack(side=tk.LEFT, padx=(10, 3))
        ttk.Entry(top, textvariable=self.max_points_var, width=8).pack(side=tk.LEFT)

        ttk.Label(top, text="Line Width:").pack(side=tk.LEFT, padx=(10, 3))
        ttk.Entry(top, textvariable=self.line_width_var, width=5).pack(side=tk.LEFT)

        options = ttk.Frame(self.frame, padding=(6, 0, 6, 4))
        options.pack(fill=tk.X)

        ttk.Checkbutton(options, text="Grid", variable=self.show_grid_var).pack(side=tk.LEFT)
        ttk.Checkbutton(options, text="Dark Theme", variable=self.dark_mode_var).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Checkbutton(options, text="3σ Anomalies", variable=self.anomaly_var).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Label(options, text="σ:").pack(side=tk.LEFT, padx=(4, 2))
        ttk.Entry(options, textvariable=self.sigma_var, width=4).pack(side=tk.LEFT)
        ttk.Checkbutton(options, text="Earthquakes", variable=self.show_quakes_var).pack(side=tk.LEFT, padx=(12, 0))
        ttk.Checkbutton(options, text="Solar Flares", variable=self.show_flares_var).pack(side=tk.LEFT, padx=(8, 0))

        body = ttk.Panedwindow(self.frame, orient=tk.HORIZONTAL)
        body.pack(fill=tk.BOTH, expand=True, padx=6, pady=(0, 4))

        left = ttk.Frame(body, padding=4)
        self.plot_frame = ttk.Frame(body)
        body.add(left, weight=1)
        body.add(self.plot_frame, weight=6)

        ttk.Label(left, text="X Axis:").pack(anchor=tk.W)
        self.x_combo = ttk.Combobox(left, state="readonly")
        self.x_combo.pack(fill=tk.X, pady=(2, 7))

        ttk.Label(left, text="Channels:").pack(anchor=tk.W)
        self.channel_frame = ScrollableCheckFrame(left)
        self.channel_frame.pack(fill=tk.BOTH, expand=True, pady=(2, 5))

        ttk.Button(left, text="Select All", command=self.select_all).pack(fill=tk.X)
        ttk.Button(left, text="Clear Selection", command=self.clear_selection).pack(fill=tk.X, pady=(3, 0))
        ttk.Button(left, text="DHO/HWU Only", command=self.select_main_vlf).pack(fill=tk.X, pady=(3, 0))

        bottom = ttk.Frame(self.frame, padding=(6, 2, 6, 6))
        bottom.pack(fill=tk.X)
        self.progress = ttk.Progressbar(bottom, mode="indeterminate", length=150)
        self.progress.pack(side=tk.LEFT)
        ttk.Label(bottom, textvariable=self.status_var).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Label(bottom, textvariable=self.cursor_var).pack(side=tk.RIGHT)

    def _new_figure(self, rows=1):
        for widget in self.plot_frame.winfo_children():
            widget.destroy()

        self.figure = Figure(figsize=(11, max(5.5, rows * 2.1)), dpi=100)
        self.canvas = FigureCanvasTkAgg(self.figure, master=self.plot_frame)
        self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        toolbar = NavigationToolbar2Tk(self.canvas, self.plot_frame, pack_toolbar=False)
        toolbar.update()
        toolbar.pack(fill=tk.X)

        self.canvas.mpl_connect("motion_notify_event", self.on_mouse_move)
        self.axes = []
        self.lines = []

    def set_busy(self, busy, text):
        self.loading = busy
        self.status_var.set(text)
        if busy:
            self.progress.start(10)
        else:
            self.progress.stop()

    def load_async(self):
        self.set_busy(True, "Loading file...")
        threading.Thread(target=self._load_worker, daemon=True).start()

    def _detect_text_format(self):
        with open(self.file_path, "rb") as f:
            raw = f.read(128 * 1024)

        sample, encoding = None, None
        for enc in ("utf-8-sig", "utf-8", "cp1250", "latin1"):
            try:
                sample = raw.decode(enc)
                encoding = enc
                break
            except UnicodeDecodeError:
                pass

        if sample is None:
            raise ValueError("Nepodarilo sa rozpoznať kódovanie.")

        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            sep = dialect.delimiter
        except csv.Error:
            first = next((x for x in sample.splitlines() if x.strip()), "")
            counts = {d: first.count(d) for d in (",", ";", "\t", "|")}
            sep = max(counts, key=counts.get)
            if counts[sep] == 0:
                raise ValueError("Nepodarilo sa rozpoznať oddeľovač.")

        return sep, encoding

    def _load_worker(self):
        try:
            ext = os.path.splitext(self.file_path)[1].lower()

            if ext in (".xlsx", ".xlsm"):
                df = pd.read_excel(self.file_path, engine="openpyxl")
            elif ext == ".xls":
                df = pd.read_excel(self.file_path, engine="xlrd")
            elif ext in (".csv", ".txt"):
                sep, encoding = self._detect_text_format()
                df = pd.read_csv(
                    self.file_path,
                    sep=sep,
                    engine="c",
                    low_memory=True,
                    memory_map=True,
                    encoding=encoding,
                )
            else:
                raise ValueError("Nepodporovaný typ súboru.")

            if df.empty:
                raise ValueError("Súbor neobsahuje údaje.")

            df = df.dropna(axis=1, how="all")
            df.columns = [str(c).strip() for c in df.columns]
            self.df = df
            self.app.root.after(0, self._after_load)

        except Exception as exc:
            self.app.root.after(0, lambda e=exc: self._load_error(e))

    def _after_load(self):
        columns = list(self.df.columns)
        numeric = []

        for col in columns:
            if pd.api.types.is_numeric_dtype(self.df[col]):
                numeric.append(col)
            else:
                converted = pd.to_numeric(
                    self.df[col].astype(str).str.replace(",", ".", regex=False),
                    errors="coerce",
                )
                if converted.notna().mean() >= 0.80:
                    self.df[col] = converted
                    numeric.append(col)

        self.numeric_columns = numeric
        time_candidates = [
            c for c in columns
            if any(k in c.lower() for k in ("time", "čas", "date", "datum", "dátum", "hh:mm"))
        ]
        self.x_combo["values"] = ["Row Number"] + columns
        self.x_combo.set(time_candidates[0] if time_candidates else "Row Number")

        for widget in self.channel_frame.inner.winfo_children():
            widget.destroy()

        self.channel_vars.clear()
        selected = 0
        for col in numeric:
            var = tk.BooleanVar(value=False)
            self.channel_vars[col] = var
            ttk.Checkbutton(self.channel_frame.inner, text=col, variable=var).pack(
                anchor=tk.W, fill=tk.X
            )
            if col != self.x_combo.get() and selected < 5:
                var.set(True)
                selected += 1

        size_mb = os.path.getsize(self.file_path) / 1024**2
        self.set_busy(
            False,
            f"Načítané: {len(self.df):,} riadkov, {len(columns)} stĺpcov, {size_mb:.1f} MB.",
        )
        self.plot_selected()

    def _load_error(self, exc):
        self.set_busy(False, "Loading failed.")
        messagebox.showerror(APP_TITLE, f"{type(exc).__name__}: {exc}")

    def _date_from_filename(self):
        name = os.path.basename(self.file_path)

        for match in re.finditer(r"(?<!\d)(\d{2})(\d{2})(\d{4})(?!\d)", name):
            d, m, y = map(int, match.groups())
            try:
                return pd.Timestamp(y, m, d)
            except ValueError:
                pass

        for match in re.finditer(r"(?<!\d)(\d{4})(\d{2})(\d{2})(?!\d)", name):
            y, m, d = map(int, match.groups())
            try:
                return pd.Timestamp(y, m, d)
            except ValueError:
                pass

        return None

    def _find_date_column(self, excluded):
        preferred, others = [], []

        for col in self.df.columns:
            if col == excluded:
                continue
            if any(k in col.lower() for k in ("date", "datum", "dátum", "day", "deň")):
                preferred.append(col)
            else:
                others.append(col)

        for col in preferred + others:
            s = self.df[col]
            if pd.api.types.is_numeric_dtype(s) and col not in preferred:
                continue

            parsed = pd.to_datetime(s, errors="coerce", dayfirst=True)
            if parsed.notna().mean() >= 0.80:
                years = parsed.dropna().dt.year
                if not years.empty and years.between(1990, 2100).mean() >= 0.90:
                    return col, parsed

        return None, None

    def get_x_data(self):
        selected = self.x_combo.get()
        n = len(self.df)

        if selected == "Row Number" or selected not in self.df.columns:
            return np.arange(n), "Row Number", False

        s = self.df[selected]

        if pd.api.types.is_datetime64_any_dtype(s):
            return pd.to_datetime(s).to_numpy(), "Date & Time", True

        text = s.astype(str).str.strip()
        time_ratio = text.str.match(r"^\d{1,2}:\d{2}(:\d{2}([.,]\d+)?)?$").mean()

        if time_ratio > 0.70:
            td = pd.to_timedelta(text.str.replace(",", ".", regex=False), errors="coerce")
            if td.notna().any():
                date_col, date_series = self._find_date_column(selected)

                if date_series is not None:
                    combined = date_series.dt.normalize() + td
                    return combined.to_numpy(), f"{date_col} + {selected}", True

                seconds = td.dt.total_seconds().to_numpy(dtype=float)
                valid = pd.Series(seconds).interpolate(limit_direction="both").to_numpy()
                diffs = np.diff(valid, prepend=valid[0])
                days = np.cumsum(diffs < -(12 * 3600))

                start = self._date_from_filename() or pd.Timestamp("2000-01-01")
                combined = (
                    start
                    + pd.to_timedelta(days, unit="D")
                    + pd.to_timedelta(seconds, unit="s")
                )
                return np.asarray(combined), "Date & Time", True

        parsed = pd.to_datetime(s, errors="coerce", dayfirst=True)
        if parsed.notna().mean() > 0.70:
            return parsed.to_numpy(), "Date & Time", True

        num = pd.to_numeric(text.str.replace(",", ".", regex=False), errors="coerce")
        if num.notna().mean() > 0.70:
            return num.to_numpy(), selected, False

        return np.arange(n), "Row Number", False

    def selected_columns(self):
        return [c for c, var in self.channel_vars.items() if var.get()]

    def _plot_events(self, axes):
        if not axes or not self.x_is_datetime:
            return

        if self.show_quakes_var.get():
            for event in self.earthquakes:
                t = pd.Timestamp(event["time"])
                for ax in axes:
                    ax.axvline(t, linestyle="--", linewidth=0.8, alpha=0.65)
                axes[0].annotate(
                    f"M{event['mag']:.1f}",
                    xy=(t, 1),
                    xycoords=("data", "axes fraction"),
                    xytext=(2, -3),
                    textcoords="offset points",
                    rotation=90,
                    va="top",
                    fontsize=7,
                )

        if self.show_flares_var.get():
            for event in self.flares:
                t = pd.Timestamp(event["time"])
                for ax in axes:
                    ax.axvline(t, linestyle=":", linewidth=1.0, alpha=0.75)
                axes[0].annotate(
                    event["class"],
                    xy=(t, 1),
                    xycoords=("data", "axes fraction"),
                    xytext=(2, -3),
                    textcoords="offset points",
                    rotation=90,
                    va="top",
                    fontsize=7,
                )

    def plot_selected(self):
        if self.df is None:
            return

        selected = self.selected_columns()
        if not selected:
            messagebox.showinfo(APP_TITLE, "Please select at least one channel.")
            return

        try:
            max_points = min(int(self.max_points_var.get()), MAX_ALLOWED_POINTS)
            line_width = float(self.line_width_var.get().replace(",", "."))
            sigma = float(self.sigma_var.get().replace(",", "."))
            if max_points < 100 or line_width <= 0 or sigma <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror(APP_TITLE, "Skontroluj počet bodov, hrúbku a hodnotu σ.")
            return

        rows = len(selected) if self.plot_mode_var.get() == "Stacked Panels" else 1
        self._new_figure(rows)

        if rows == 1:
            axes = [self.figure.add_subplot(111)]
        else:
            axes_obj = self.figure.subplots(rows, 1, sharex=True)
            axes = list(np.atleast_1d(axes_obj))

        self.axes = axes
        x, xlabel, is_dt = self.get_x_data()
        self.x_full = np.asarray(x)
        self.x_is_datetime = is_dt
        self.last_selected_columns = selected

        dark = self.dark_mode_var.get()
        bg = "#111214" if dark else "white"
        fg = "white" if dark else "black"
        self.figure.patch.set_facecolor("#202124" if dark else "white")

        for i, col in enumerate(selected):
            ax = axes[i] if rows > 1 else axes[0]
            y = pd.to_numeric(self.df[col], errors="coerce").to_numpy(dtype=float)
            valid = pd.notna(self.x_full) & np.isfinite(y)

            if not valid.any():
                continue

            xv, yv = self.x_full[valid], y[valid]
            xp, yp = reduce_minmax(xv, yv, max_points)
            line, = ax.plot(xp, yp, linewidth=line_width, label=col)
            self.lines.append((col, line))

            if self.anomaly_var.get():
                mean = np.nanmean(yv)
                std = np.nanstd(yv)
                mask = np.abs(yv - mean) > sigma * std
                anomaly_idx = np.flatnonzero(mask)
                if len(anomaly_idx) > 3000:
                    anomaly_idx = anomaly_idx[::math.ceil(len(anomaly_idx) / 3000)]
                ax.scatter(xv[anomaly_idx], yv[anomaly_idx], s=9, marker="x", label=f"{col} > {sigma:g}σ")

            ax.set_facecolor(bg)
            ax.tick_params(colors=fg)
            ax.grid(self.show_grid_var.get(), alpha=0.3)
            for spine in ax.spines.values():
                spine.set_color(fg)

            if rows > 1:
                ax.set_ylabel(col, color=fg)
            else:
                ax.set_ylabel("Hodnota", color=fg)

        axes[0].set_title(os.path.basename(self.file_path), color=fg)
        axes[-1].set_xlabel(xlabel, color=fg)

        if rows == 1:
            axes[0].legend(loc="best", fontsize=8)

        if is_dt:
            axes[-1].xaxis.set_major_locator(mdates.HourLocator(byhour=[0, 12]))
            axes[-1].xaxis.set_major_formatter(mdates.DateFormatter("%d.%m.%Y\n%H:%M"))

        self._plot_events(axes)
        self.figure.tight_layout()
        self.canvas.draw_idle()
        self.status_var.set(f"Plot complete: {len(selected)} kanálov.")

    def _datetime_series(self):
        x, _, is_dt = self.get_x_data()
        if not is_dt:
            raise ValueError("Denné funkcie vyžadujú dátumovú/časovú os.")
        return pd.to_datetime(x)

    def compute_daily_stats(self):
        dt = self._datetime_series()
        selected = self.selected_columns()
        if not selected:
            raise ValueError("Nie je vybraný žiadny kanál.")

        work = pd.DataFrame({"datetime": dt})
        for col in selected:
            work[col] = pd.to_numeric(self.df[col], errors="coerce").to_numpy()

        work["date"] = work["datetime"].dt.date
        stats = work.groupby("date")[selected].agg(["mean", "min", "max", "std", "count"])
        self.daily_stats = stats
        return stats

    def show_daily_stats(self):
        try:
            stats = self.compute_daily_stats()
        except Exception as exc:
            messagebox.showerror(APP_TITLE, str(exc))
            return

        win = tk.Toplevel(self.frame)
        win.title("Daily Statistics")
        win.geometry("1050x600")

        text = tk.Text(win, wrap="none", font=("Consolas", 9))
        ybar = ttk.Scrollbar(win, orient=tk.VERTICAL, command=text.yview)
        xbar = ttk.Scrollbar(win, orient=tk.HORIZONTAL, command=text.xview)
        text.configure(yscrollcommand=ybar.set, xscrollcommand=xbar.set)

        text.grid(row=0, column=0, sticky="nsew")
        ybar.grid(row=0, column=1, sticky="ns")
        xbar.grid(row=1, column=0, sticky="ew")
        win.rowconfigure(0, weight=1)
        win.columnconfigure(0, weight=1)

        text.insert("1.0", stats.round(4).to_string())
        text.configure(state="disabled")

        def save_stats():
            path = filedialog.asksaveasfilename(
                parent=win,
                defaultextension=".csv",
                filetypes=[("CSV", "*.csv"), ("Excel", "*.xlsx")],
            )
            if not path:
                return
            if path.lower().endswith(".xlsx"):
                stats.to_excel(path)
            else:
                stats.to_csv(path, sep=";", decimal=",")
            messagebox.showinfo(APP_TITLE, f"Štatistiky uložené:\n{path}", parent=win)

        ttk.Button(win, text="Uložiť štatistiky", command=save_stats).grid(
            row=2, column=0, sticky="w", padx=5, pady=5
        )

    def compare_days(self):
        try:
            dt = self._datetime_series()
        except Exception as exc:
            messagebox.showerror(APP_TITLE, str(exc))
            return

        selected = self.selected_columns()
        if not selected:
            messagebox.showinfo(APP_TITLE, "Please select a channel.")
            return

        dates = sorted(pd.Series(dt.date).dropna().unique())
        if len(dates) < 2:
            messagebox.showinfo(APP_TITLE, "Súbor musí obsahovať aspoň dva dni.")
            return

        win = tk.Toplevel(self.frame)
        win.title("Porovnanie dvoch dní")
        win.geometry("420x230")

        channel_var = tk.StringVar(value=selected[0])
        date1_var = tk.StringVar(value=str(dates[0]))
        date2_var = tk.StringVar(value=str(dates[1]))

        ttk.Label(win, text="Kanál:").pack(anchor=tk.W, padx=12, pady=(12, 2))
        ttk.Combobox(win, textvariable=channel_var, values=self.numeric_columns, state="readonly").pack(
            fill=tk.X, padx=12
        )
        ttk.Label(win, text="Prvý deň:").pack(anchor=tk.W, padx=12, pady=(8, 2))
        ttk.Combobox(win, textvariable=date1_var, values=[str(d) for d in dates], state="readonly").pack(
            fill=tk.X, padx=12
        )
        ttk.Label(win, text="Druhý deň:").pack(anchor=tk.W, padx=12, pady=(8, 2))
        ttk.Combobox(win, textvariable=date2_var, values=[str(d) for d in dates], state="readonly").pack(
            fill=tk.X, padx=12
        )

        def create_comparison():
            col = channel_var.get()
            d1 = pd.Timestamp(date1_var.get()).date()
            d2 = pd.Timestamp(date2_var.get()).date()
            y = pd.to_numeric(self.df[col], errors="coerce").to_numpy(dtype=float)

            fig_win = tk.Toplevel(self.frame)
            fig_win.title(f"Porovnanie dní – {col}")
            fig_win.geometry("1100x750")

            fig = Figure(figsize=(10, 7), dpi=100)
            axes = fig.subplots(2, 1, sharex=True)

            for ax, day in zip(axes, (d1, d2)):
                mask = np.array(dt.date == day)
                times = dt[mask]
                values = y[mask]
                seconds = (
                    times.hour * 3600
                    + times.minute * 60
                    + times.second
                    + times.microsecond / 1e6
                )
                ax.plot(seconds / 3600, values, linewidth=0.8)
                ax.set_title(str(day))
                ax.set_ylabel(col)
                ax.grid(True, alpha=0.3)

            axes[-1].set_xlabel("Hodina dňa")
            axes[-1].set_xlim(0, 24)
            axes[-1].set_xticks(np.arange(0, 25, 2))
            fig.tight_layout()

            canvas = FigureCanvasTkAgg(fig, master=fig_win)
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
            nav = NavigationToolbar2Tk(canvas, fig_win, pack_toolbar=False)
            nav.update()
            nav.pack(fill=tk.X)
            canvas.draw()

        ttk.Button(win, text="Plot porovnanie", command=create_comparison).pack(pady=12)

    def _data_time_range_utc(self):
        dt = self._datetime_series()
        start = pd.Timestamp(dt.min())
        end = pd.Timestamp(dt.max())

        if start.tzinfo is None:
            start = start.tz_localize("Europe/Bratislava").tz_convert("UTC")
            end = end.tz_localize("Europe/Bratislava").tz_convert("UTC")
        else:
            start = start.tz_convert("UTC")
            end = end.tz_convert("UTC")

        return start, end

    def import_earthquakes(self):
        try:
            start, end = self._data_time_range_utc()
        except Exception as exc:
            messagebox.showerror(APP_TITLE, str(exc))
            return

        mag = simpledialog.askfloat(
            "Import USGS Earthquakes",
            "Minimálna magnitúda:",
            initialvalue=4.5,
            minvalue=0.0,
            maxvalue=10.0,
            parent=self.frame,
        )
        if mag is None:
            return

        self.set_busy(True, "Načítavam zemetrasenia z USGS...")

        def worker():
            try:
                params = {
                    "format": "geojson",
                    "starttime": start.strftime("%Y-%m-%dT%H:%M:%S"),
                    "endtime": end.strftime("%Y-%m-%dT%H:%M:%S"),
                    "minmagnitude": str(mag),
                    "orderby": "time-asc",
                    "limit": "20000",
                }
                url = USGS_QUERY_URL + "?" + urllib.parse.urlencode(params)
                req = urllib.request.Request(url, headers={"User-Agent": "VLFDataGraphPro/1.0"})
                with urllib.request.urlopen(req, timeout=40) as response:
                    data = json.load(response)

                events = []
                for feature in data.get("features", []):
                    props = feature.get("properties", {})
                    ms = props.get("time")
                    if ms is None:
                        continue
                    utc_time = pd.to_datetime(ms, unit="ms", utc=True)
                    local_time = utc_time.tz_convert("Europe/Bratislava").tz_localize(None)
                    events.append({
                        "time": local_time,
                        "mag": safe_float(props.get("mag")),
                        "place": props.get("place") or "",
                        "url": props.get("url") or "",
                    })

                self.earthquakes = events
                self.app.root.after(
                    0,
                    lambda: (
                        self.set_busy(False, f"USGS: načítaných {len(events)} udalostí."),
                        self.plot_selected(),
                    ),
                )
            except Exception as exc:
                self.app.root.after(
                    0,
                    lambda e=exc: (
                        self.set_busy(False, "USGS import zlyhal."),
                        messagebox.showerror(APP_TITLE, f"USGS import:\n{e}"),
                    ),
                )

        threading.Thread(target=worker, daemon=True).start()

    def import_flares(self):
        try:
            start, end = self._data_time_range_utc()
        except Exception as exc:
            messagebox.showerror(APP_TITLE, str(exc))
            return

        self.set_busy(True, "Načítavam NOAA GOES röntgenový tok...")

        def worker():
            try:
                req = urllib.request.Request(
                    NOAA_GOES_XRAY_URL,
                    headers={"User-Agent": "VLFDataGraphPro/1.0"},
                )
                with urllib.request.urlopen(req, timeout=40) as response:
                    raw = json.load(response)

                rows = []
                for item in raw:
                    energy = str(item.get("energy", ""))
                    if "0.1-0.8" not in energy:
                        continue
                    t = pd.to_datetime(item.get("time_tag"), utc=True, errors="coerce")
                    flux = safe_float(item.get("flux"))
                    if pd.isna(t) or not np.isfinite(flux):
                        continue
                    rows.append((t, flux))

                if not rows:
                    raise ValueError("NOAA odpoveď neobsahovala očakávaný GOES kanál 0.1–0.8 nm.")

                goes = pd.DataFrame(rows, columns=["time", "flux"]).sort_values("time")
                goes = goes[(goes["time"] >= start) & (goes["time"] <= end)]

                # Peaks above C1.0. Minimum separation 10 minutes.
                flux = goes["flux"].to_numpy()
                times = goes["time"].to_numpy()
                candidate = np.flatnonzero(
                    (flux >= 1e-6)
                    & (flux >= np.roll(flux, 1))
                    & (flux > np.roll(flux, -1))
                )

                peaks = []
                last_time = None
                for idx in candidate:
                    t = pd.Timestamp(times[idx])
                    if last_time is not None and (t - last_time).total_seconds() < 600:
                        if peaks and flux[idx] > peaks[-1]["flux"]:
                            peaks[-1] = {
                                "time": t.tz_convert("Europe/Bratislava").tz_localize(None),
                                "flux": flux[idx],
                                "class": flare_class(flux[idx]),
                            }
                        continue

                    peaks.append({
                        "time": t.tz_convert("Europe/Bratislava").tz_localize(None),
                        "flux": flux[idx],
                        "class": flare_class(flux[idx]),
                    })
                    last_time = t

                self.flares = peaks
                self.app.root.after(
                    0,
                    lambda: (
                        self.set_busy(False, f"NOAA/GOES: nájdených {len(peaks)} erupcií ≥ C1."),
                        self.plot_selected(),
                    ),
                )
            except Exception as exc:
                self.app.root.after(
                    0,
                    lambda e=exc: (
                        self.set_busy(False, "NOAA import zlyhal."),
                        messagebox.showerror(
                            APP_TITLE,
                            "NOAA/GOES import zlyhal.\n"
                            "Online zdroj poskytuje najmä posledných 7 dní.\n\n"
                            f"{e}",
                        ),
                    ),
                )

        threading.Thread(target=worker, daemon=True).start()

    def save_graph(self):
        if not self.axes:
            return

        default = os.path.splitext(os.path.basename(self.file_path))[0] + "_graf.png"
        path = filedialog.asksaveasfilename(
            defaultextension=".png",
            initialfile=default,
            filetypes=[("PNG", "*.png"), ("PDF", "*.pdf"), ("SVG", "*.svg")],
        )
        if path:
            self.figure.savefig(path, dpi=300, bbox_inches="tight")
            self.status_var.set(f"Graph saved: {path}")

    def export_protocol(self):
        if self.df is None or not self.axes:
            messagebox.showinfo(APP_TITLE, "Najprv vytvor graf.")
            return

        try:
            from docx import Document
            from docx.shared import Cm
        except ImportError:
            messagebox.showerror(
                APP_TITLE,
                "Pre export DOCX nainštaluj:\n\npython -m pip install python-docx",
            )
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=os.path.splitext(os.path.basename(self.file_path))[0] + "_protokol.docx",
            filetypes=[("Word dokument", "*.docx")],
        )
        if not path:
            return

        image_path = os.path.splitext(path)[0] + "_temp_graph.png"

        try:
            self.figure.savefig(image_path, dpi=220, bbox_inches="tight")
            stats = self.compute_daily_stats()

            doc = Document()
            doc.add_heading("VLF DataGraph – Analysis Report", level=1)
            doc.add_paragraph(f"Source file: {os.path.basename(self.file_path)}")

            dt = self._datetime_series()
            doc.add_paragraph(
                f"Time range: {dt.min().strftime('%d.%m.%Y %H:%M:%S')} – "
                f"{dt.max().strftime('%d.%m.%Y %H:%M:%S')}"
            )
            doc.add_paragraph("Selected channels: " + ", ".join(self.selected_columns()))
            doc.add_paragraph(f"Number of samples: {len(self.df):,}".replace(",", " "))

            if self.anomaly_var.get():
                doc.add_paragraph(f"V grafe sú vyznačené anomálie nad {self.sigma_var.get()}σ.")

            doc.add_heading("Graphical Results", level=2)
            doc.add_picture(image_path, width=Cm(16.5))

            doc.add_heading("Daily Statistics", level=2)
            flat = stats.reset_index()
            flat.columns = [
                str(c[0]) if not isinstance(c, tuple) or not c[1] else f"{c[0]} – {c[1]}"
                for c in flat.columns
            ]

            table = doc.add_table(rows=1, cols=len(flat.columns))
            table.style = "Table Grid"
            for i, col in enumerate(flat.columns):
                table.rows[0].cells[i].text = str(col)

            for _, row in flat.iterrows():
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    if isinstance(value, (float, np.floating)):
                        cells[i].text = f"{value:.4f}"
                    else:
                        cells[i].text = str(value)

            if self.earthquakes:
                doc.add_heading("Earthquakes USGS", level=2)
                for e in self.earthquakes[:100]:
                    doc.add_paragraph(
                        f"{pd.Timestamp(e['time']).strftime('%d.%m.%Y %H:%M:%S')} – "
                        f"M{e['mag']:.1f}, {e['place']}"
                    )

            if self.flares:
                doc.add_heading("NOAA GOES Solar Flares", level=2)
                for e in self.flares:
                    doc.add_paragraph(
                        f"{pd.Timestamp(e['time']).strftime('%d.%m.%Y %H:%M:%S')} – "
                        f"{e['class']}, tok {e['flux']:.3e} W/m²"
                    )

            doc.add_paragraph(
                "Note: označenie 3σ predstavuje štatistickú odchýlku od priemeru "
                "vybraného kanála a samo osebe neurčuje fyzikálnu príčinu udalosti."
            )
            doc.save(path)
            messagebox.showinfo(APP_TITLE, f"Protokol uložený:\n{path}")
        except Exception as exc:
            messagebox.showerror(APP_TITLE, f"Export Report zlyhal:\n{exc}")
        finally:
            try:
                if os.path.exists(image_path):
                    os.remove(image_path)
            except OSError:
                pass

    def on_mouse_move(self, event):
        if (
            self.df is None
            or event.inaxes not in self.axes
            or event.xdata is None
            or self.x_full is None
        ):
            return

        try:
            if self.x_is_datetime:
                target = pd.Timestamp(mdates.num2date(event.xdata)).tz_localize(None)
                arr = pd.to_datetime(self.x_full).to_numpy()
                idx = int(np.argmin(np.abs(arr - np.datetime64(target))))
                label = pd.Timestamp(arr[idx]).strftime("%d.%m.%Y %H:%M:%S")
            else:
                arr = pd.to_numeric(pd.Series(self.x_full), errors="coerce").to_numpy(float)
                idx = int(np.nanargmin(np.abs(arr - event.xdata)))
                label = f"X={arr[idx]:.3f}"

            vals = []
            for col in self.last_selected_columns[:6]:
                val = safe_float(self.df[col].iloc[idx])
                if np.isfinite(val):
                    vals.append(f"{col}={val:.2f}")

            self.cursor_var.set(label + " | " + " | ".join(vals))
        except Exception:
            pass

    def select_all(self):
        for var in self.channel_vars.values():
            var.set(True)

    def clear_selection(self):
        for var in self.channel_vars.values():
            var.set(False)

    def select_main_vlf(self):
        for col, var in self.channel_vars.items():
            name = col.lower()
            var.set("dho" in name or "hwu" in name)


class VLFDataGraphSessionsApp:
    def __init__(self, root):
        self.root = root
        self.root.title(f"{APP_TITLE} {APP_VERSION}")
        self.root.geometry("1600x930")
        self.root.minsize(1100, 700)

        self._build_menu()

        top = ttk.Frame(root, padding=6)
        top.pack(fill=tk.X)

        ttk.Button(top, text="New Session / Open File", command=self.open_files).pack(side=tk.LEFT)
        ttk.Button(top, text="Close Session", command=self.close_current).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Label(
            top,
            text="Each opened file has its own tab, settings, events and plot.",
        ).pack(side=tk.LEFT, padx=(18, 0))

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        self.sessions = []
        self.open_files()

    def _build_menu(self):
        menubar = tk.Menu(self.root)

        file_menu = tk.Menu(menubar, tearoff=False)
        file_menu.add_command(label="New Session / Open File", command=self.open_files)
        file_menu.add_command(label="Close Session", command=self.close_current)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.destroy)
        menubar.add_cascade(label="File", menu=file_menu)

        help_menu = tk.Menu(menubar, tearoff=False)
        help_menu.add_command(label="User Guide", command=self.show_user_guide)
        help_menu.add_command(label="Check for Updates", command=self.check_for_updates)
        help_menu.add_separator()
        help_menu.add_command(label="About", command=self.show_about)
        menubar.add_cascade(label="Help", menu=help_menu)

        self.root.config(menu=menubar)

    def show_about(self):
        messagebox.showinfo(
            "About VLF DataGraph Pro",
            f"{APP_TITLE}\n"
            f"Version {APP_VERSION}\n\n"
            "VLF/LF time-series visualization and analysis tool.\n\n"
            f"© 2026 {APP_AUTHOR}\n"
            f"{APP_LICENSE}"
        )

    def show_user_guide(self):
        guide = (
            "VLF DataGraph Pro – Quick User Guide\n\n"
            "1. Open one or more CSV, TXT, XLS, XLSX or XLSM files.\n"
            "2. Each recording opens in a separate Session tab.\n"
            "3. Select the X axis and one or more signal channels.\n"
            "4. Choose Single Plot or Stacked Panels and click Plot.\n"
            "5. Enable 3σ Anomalies to mark statistically unusual values.\n"
            "6. Use Compare Days for two-day channel comparison.\n"
            "7. Use Daily Statistics for mean, minimum, maximum, standard deviation and sample count.\n"
            "8. Import USGS earthquakes or NOAA GOES solar flares to the time axis.\n"
            "9. Save the plot as PNG, PDF or SVG.\n"
            "10. Export Report creates a DOCX report with plot and statistics.\n\n"
            "Notes:\n"
            "• NOAA GOES online data are mainly available for the most recent seven days.\n"
            "• A 3σ marker is a statistical flag and does not determine the physical cause.\n"
            "• For DOCX export, install python-docx."
        )

        win = tk.Toplevel(self.root)
        win.title("User Guide")
        win.geometry("760x600")

        text_widget = tk.Text(win, wrap="word", padx=15, pady=15)
        text_widget.insert("1.0", guide)
        text_widget.configure(state="disabled")
        text_widget.pack(fill=tk.BOTH, expand=True)

        ttk.Button(win, text="Close", command=win.destroy).pack(pady=(0, 10))

    def check_for_updates(self):
        messagebox.showinfo(
            "Check for Updates",
            f"You are using {APP_TITLE} {APP_VERSION}.\n\n"
            "Automatic online update checking is not configured in this experimental release."
        )

    def open_files(self):
        paths = filedialog.askopenfilenames(
            title="Select one or more recordings",
            filetypes=[
                ("Dátové súbory", "*.csv *.txt *.xls *.xlsx *.xlsm"),
                ("CSV/TXT", "*.csv *.txt"),
                ("Excel", "*.xls *.xlsx *.xlsm"),
                ("Všetky súbory", "*.*"),
            ],
        )

        for path in paths:
            session = SessionTab(self, self.notebook, path)
            self.sessions.append(session)

    def close_current(self):
        current = self.notebook.select()
        if not current:
            return

        for session in list(self.sessions):
            if str(session.frame) == current:
                self.notebook.forget(session.frame)
                self.sessions.remove(session)
                break


def main():
    root = tk.Tk()
    VLFDataGraphSessionsApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
